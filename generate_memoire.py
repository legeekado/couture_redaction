#!/usr/bin/env python3
"""Génère le mémoire Kayñiawlu au format .docx (guide ISI Master 2 GL, juin 2026)."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent / (
    "Conception et réalisation d'une solution SaaS de gestion pour les couturiers "
    "du Sénégal - Kayniawlu.docx"
)

# --- Personnalisation rapide (modifier ici) ---
ETUDIANT = "[VOTRE PRÉNOM ET NOM]"
ENCADREUR = "[NOM DE VOTRE ENCADREUR]"
TITRE_ENCADREUR = "[Fonction / entreprise de l'encadreur]"
ANNEE_ACADEMIQUE = "2025 - 2026"
TITRE_MEMOIRE = (
    "Conception et réalisation d'une solution SaaS de gestion "
    "pour les couturiers du Sénégal : cas de Kayñiawlu"
)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level, size in [(1, 16), (2, 14), (3, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld)
    note = doc.add_paragraph(
        "Note : ouvrir le document dans Word, clic droit sur la table des matières "
        "→ « Mettre à jour les champs » → « Mettre à jour toute la table »."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(10)


def center(doc: Document, text: str, *, bold=False, size=12, space_after=6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)


def body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def table_from_rows(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph("")


def add_cover(doc: Document) -> None:
    center(doc, "REPUBLIQUE DU SENEGAL", bold=True, size=13)
    center(doc, "Un peuple – Un but – Une foi", size=12, space_after=12)
    center(
        doc,
        "Ministère de l'Enseignement supérieur, de la Recherche et de l'Innovation\n"
        "Direction de l'Enseignement Supérieur Privé\n"
        "Institut Supérieur d'Informatique",
        size=12,
        space_after=24,
    )
    center(doc, "ISI", bold=True, size=14, space_after=18)
    center(
        doc,
        "Mémoire de fin de cycle pour l'obtention du diplôme de\n"
        "Master en Génie Logiciel",
        size=12,
        space_after=24,
    )
    center(doc, TITRE_MEMOIRE, bold=True, size=14, space_after=36)
    center(doc, "Présenté et soutenu par :", size=12)
    center(doc, ETUDIANT, bold=True, size=13, space_after=24)
    center(doc, "Sous la direction de :", size=12)
    center(doc, ENCADREUR, bold=True, size=13)
    center(doc, TITRE_ENCADREUR, size=12, space_after=36)
    center(doc, f"Année Académique : {ANNEE_ACADEMIQUE}", size=12)
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("Dédicace", level=1)
    body(
        doc,
        "[À personnaliser] Je dédie ce mémoire à mes parents, à ma famille et à toutes "
        "les personnes qui m'ont soutenu(e) tout au long de mon parcours académique.",
    )
    doc.add_page_break()

    doc.add_heading("Remerciements", level=1)
    body(
        doc,
        "Je remercie Dieu pour la force et la persévérance accordées durant ce travail. "
        f"J'exprime ma profonde gratitude à mon encadreur {ENCADREUR} pour sa disponibilité, "
        "ses conseils et son accompagnement. Je remercie également l'ensemble du corps "
        "professoral de l'ISI, mes camarades de promotion, ainsi que toutes les personnes "
        "qui, de près ou de loin, ont contribué à la réalisation de ce mémoire.",
    )
    doc.add_page_break()

    doc.add_heading("Avant-propos", level=1)
    body(
        doc,
        "Fondé en 1988, le Groupe ISI (Institut Supérieur d'Informatique) forme depuis "
        "plus de trois décennies des professionnels de l'informatique, des télécommunications "
        "et de la gestion. L'obtention du Master en Génie Logiciel est sanctionnée par la "
        "rédaction et la soutenance d'un mémoire de fin de cycle."
    )
    body(
        doc,
        f"Le présent document porte sur : {TITRE_MEMOIRE}. Il décrit la conception et la "
        "réalisation de Kayñiawlu, solution SaaS multi-tenant destinée aux ateliers de "
        "couture sénégalais, composée d'une API Laravel, d'une application mobile Flutter, "
        "d'un backoffice web Next.js et d'une landing page commerciale."
    )
    body(
        doc,
        "Ce mémoire étant mon premier travail de recherche appliquée, j'implore l'indulgence "
        "du jury quant aux imperfections que ce document pourrait contenir.",
    )
    doc.add_page_break()

    doc.add_heading("Sommaire", level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("Glossaire", level=1)
    glossaire = [
        ("API", "Application Programming Interface"),
        ("BPMN", "Business Process Model and Notation"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("FCFA", "Franc CFA"),
        ("FCM", "Firebase Cloud Messaging"),
        ("GraphQL", "Langage de requête pour API"),
        ("IHM", "Interface Homme-Machine"),
        ("ORM", "Object-Relational Mapping"),
        ("PDV", "Point de vente"),
        ("REST", "Representational State Transfer"),
        ("SaaS", "Software as a Service"),
        ("SGBD", "Système de Gestion de Base de Données"),
        ("UML", "Unified Modeling Language"),
        ("VAPID", "Voluntary Application Server Identification"),
    ]
    table_from_rows(doc, ["Sigle", "Signification"], glossaire)
    doc.add_page_break()

    doc.add_heading("Liste des figures", level=1)
    figures = [
        "Figure 1 : Organigramme simplifié du secteur de la couture artisanale au Sénégal",
        "Figure 2 : Diagramme BPMN du processus existant de gestion des commandes en atelier",
        "Figure 3 : Diagramme BPMN du nouveau processus supporté par Kayñiawlu",
        "Figure 4 : Diagramme de cas d'utilisation — Couturier propriétaire",
        "Figure 5 : Diagramme de cas d'utilisation — Employé d'atelier",
        "Figure 6 : Diagramme de cas d'utilisation — Super administrateur plateforme",
        "Figure 7 : Diagramme de classes du domaine métier Kayñiawlu",
        "Figure 8 : Schéma relationnel de la base de données PostgreSQL",
        "Figure 9 : Architecture fonctionnelle trois-tiers de Kayñiawlu",
        "Figure 10 : Architecture technique et déploiement Docker Compose",
        "Figure 11 : Maquette — écran d'accueil mobile",
        "Figure 12 : Maquette — dashboard backoffice atelier",
        "Figure 13 : Maquette — landing page Kayñiawlu",
    ]
    for f in figures:
        bullet(doc, f + " [À insérer]")
    doc.add_page_break()

    doc.add_heading("Liste des tableaux", level=1)
    tableaux = [
        "Tableau 1 : Structuration du contexte — processus, acteurs et difficultés",
        "Tableau 2 : Fonctionnalités de la solution par acteur",
        "Tableau 3 : Correspondance difficultés / fonctionnalités",
        "Tableau 4 : Contraintes de qualité (objectifs non fonctionnels)",
        "Tableau 5 : Résultats attendus du mémoire (périmètre du stage)",
        "Tableau 6 : Comparaison fonctionnelle des solutions similaires",
        "Tableau 7 : Comparaison non fonctionnelle des solutions similaires",
        "Tableau 8 : Description des entités principales",
        "Tableau 9 : Plans SaaS et limites par organisation",
        "Tableau 10 : Stack technique retenue",
    ]
    for t in tableaux:
        bullet(doc, t)
    doc.add_page_break()

    doc.add_heading("Résumé", level=1)
    body(
        doc,
        "Au Sénégal, la couture artisanale et semi-industrielle occupe une place centrale "
        "dans l'économie informelle et la mode locale. Les ateliers gèrent encore leurs "
        "clients, mesures et commandes via cahiers papier, fichiers Excel ou messageries "
        "instantanées, ce qui engendre pertes d'information, retards de livraison et "
        "difficultés de suivi multi-sites."
    )
    body(
        doc,
        "Ce mémoire présente la conception et la réalisation de Kayñiawlu, une solution "
        "SaaS multi-tenant composée d'une API Laravel (PostgreSQL, GraphQL, Sanctum), "
        "d'une application mobile Flutter (mode hors-ligne), d'un backoffice Next.js et "
        "d'une landing page. La plateforme digitalise le processus de gestion des commandes "
        "de couture : enregistrement client, fiches de mesures paramétrables, suivi des "
        "statuts, rappels de livraison, gestion d'équipe et abonnements par paliers."
    )
    body(
        doc,
        "Les travaux ont abouti à un prototype fonctionnel déployable en local via Docker, "
        "avec 42 tests automatisés côté backend. Les perspectives incluent le paiement "
        "mobile (Wave / Orange Money), les notifications clients (SMS / WhatsApp) et la "
        "publication sur les stores mobiles."
    )
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    body(
        doc,
        "In Senegal, tailoring workshops still rely heavily on paper notebooks, spreadsheets "
        "and instant messaging to manage customers, measurements and orders. This leads to "
        "data loss, delivery delays and poor visibility across sales outlets."
    )
    body(
        doc,
        "This thesis describes the design and implementation of Kayñiawlu, a multi-tenant SaaS "
        "platform consisting of a Laravel API (PostgreSQL, GraphQL, Sanctum), a Flutter mobile "
        "app with offline support, a Next.js back-office and a marketing landing page. "
        "The solution digitizes the tailoring order management process: customer records, "
        "configurable measurement sheets, order status tracking, delivery reminders, team "
        "management and tiered subscriptions."
    )
    body(
        doc,
        "A functional prototype can be deployed locally using Docker, with 42 automated backend "
        "tests. Future work includes mobile payment integration (Wave / Orange Money), "
        "customer notifications (SMS / WhatsApp) and app store publication."
    )
    doc.add_page_break()


def add_chapter1(doc: Document) -> None:
    doc.add_heading("Chapitre 1 : Introduction générale", level=1)

    doc.add_heading("1. Contexte", level=2)
    doc.add_heading("1.1 Présentation de l'organisation et du cadre du projet", level=3)
    body(
        doc,
        "Kayñiawlu (prononcé « kay-nya-wlu », du wolof « l'atelier de couture dans votre "
        "poche ») est un projet de startup technologique sénégalaise visant à fournir aux "
        "couturiers et couturières un outil numérique de gestion d'atelier accessible "
        "depuis un smartphone et un navigateur web."
    )
    body(
        doc,
        "Le développement s'est déroulé dans le cadre du Master 2 Génie Logiciel à l'ISI, "
        "en mode projet agile sur une période de six mois. L'équipe produit a adopté une "
        "architecture monorepo orchestrée par Docker Compose, regroupant quatre dépôts "
        "Git distincts : backend API, application mobile, backoffice web et landing page "
        "commerciale."
    )
    body(
        doc,
        "Le secteur visé regroupe les ateliers de couture sur mesure, les boutiques "
        "associées et les maisons de mode opérant à Dakar, Thiès, Saint-Louis, Touba et "
        "d'autres villes du Sénégal. Ces structures emploient généralement un propriétaire "
        "(maître couturier), des employés (coupe, couture, finition) et parfois des "
        "stagiaires. Elles servent une clientèle exigeante pour des modèles traditionnels "
        "(boubou, grand boubou, taille basse) et contemporains (costume, robe de soirée)."
    )
    body(
        doc,
        "Au Sénégal, la couture ne se limite pas à un métier artisanal : elle structure "
        "des chaînes de valeur autour du wax, de la broderie, des événements sociaux "
        "(mariages, baptêmes, Korité, Tabaski) et de la diaspora qui commande à distance. "
        "Les ateliers performants traitent plusieurs dizaines de commandes simultanées, "
        "parfois réparties entre un atelier central et des points de vente ou boutiques "
        "éphémères. Cette réalité économique exige un outil capable de centraliser l'information "
        "sans imposer une lourdeur administrative incompatible avec le rythme du métier."
    )
    body(
        doc,
        "Kayñiawlu répond à ce besoin par un modèle Software as a Service (SaaS) : "
        "chaque atelier dispose d'une organisation isolée (tenant logique), d'un plan "
        "d'abonnement (gratuit, pro, entreprise) et d'utilisateurs aux rôles différenciés. "
        "La plateforme a été conçue en cohérence avec les pratiques observées localement : "
        "montants en FCFA, numéros de téléphone sénégalais, types de vêtements courants "
        "dans les fiches de mesures, et préparation d'intégrations de paiement mobile "
        "(Wave, Orange Money) pour la monétisation future."
    )
    body(
        doc,
        "L'architecture retenue s'articule autour de quatre composants livrés dans le "
        "projet : une API backend Laravel exposant GraphQL pour les lectures et REST pour "
        "les écritures ; une application mobile Flutter pour le travail de terrain ; un "
        "backoffice web Next.js pour l'administration depuis un ordinateur ; une landing "
        "page pour la communication commerciale et l'acquisition de nouveaux ateliers."
    )

    doc.add_heading("1.2 Identification du processus existant étudié", level=3)
    body(
        doc,
        "Le processus étudié est nommé : « Processus de gestion des commandes de couture "
        "en atelier ». Il couvre la prise de contact avec le client, la saisie des mesures, "
        "la création de la commande, la production, les essayages, la livraison et le "
        "règlement."
    )
    body(
        doc,
        "Dans la majorité des ateliers observés, ce processus reste largement manuel ou "
        "partiellement informatisé (Excel, WhatsApp, photos de mesures sur téléphone)."
    )
    body(
        doc,
        "Ce processus constitue l'objet d'amélioration central du mémoire. Il est "
        "volontairement nommé et borné : il ne couvre pas la comptabilité générale ni "
        "la gestion des stocks de tissus, même si des extensions futures pourraient "
        "l'enrichir. Le focus est mis sur la relation client, la mesure, la commande et "
        "la livraison — cœur de l'activité quotidienne du couturier."
    )

    doc.add_heading("1.3 Activités ordonnées du processus existant", level=3)
    activites = [
        ("A1", "Accueillir le client et recueillir le besoin (modèle, tissu, délai)"),
        ("A2", "Prendre les mesures corporelles du client"),
        ("A3", "Enregistrer la commande (montant, acompte, date de livraison)"),
        ("A4", "Planifier la production (découpe, assemblage, finitions)"),
        ("A5", "Organiser un ou plusieurs essayages"),
        ("A6", "Finaliser la pièce et préparer la livraison"),
        ("A7", "Livrer la commande et encaisser le solde"),
        ("A8", "Archiver les informations pour de futures commandes"),
    ]
    table_from_rows(doc, ["Code", "Activité"], activites)

    doc.add_heading("1.4 Acteurs du processus existant", level=3)
    acteurs = [
        ("Client", "Personne commandant une pièce sur mesure"),
        ("Propriétaire / couturier", "Dirige l'atelier, valide prix et délais"),
        ("Employé", "Exécute les tâches de production et peut prendre les mesures"),
        ("Comptable (optionnel)", "Suit encaissements et facturation"),
    ]
    table_from_rows(doc, ["Acteur", "Rôle"], acteurs)

    doc.add_heading("1.5 Difficultés rencontrées par activité", level=3)
    body(
        doc,
        "L'analyse terrain (entretiens informels avec des ateliers de Dakar et Thiès, "
        "observation de pratiques courantes) met en évidence les difficultés suivantes. "
        "Elles constituent des contraintes fonctionnelles et non fonctionnelles à résoudre."
    )
    difficultes = [
        ("A1 — Accueil client", "Propriétaire / employé", "Doublons de fiches client ; historique dispersé entre cahiers"),
        ("A2 — Prise de mesures", "Employé", "Mesures illisibles ; pas de standard par type de vêtement ; photos perdues"),
        ("A3 — Enregistrement commande", "Propriétaire", "Oublis de montant ou d'acompte ; numérotation manuelle des commandes"),
        ("A4 — Planification", "Propriétaire", "Pas de vue globale des commandes en cours par point de vente"),
        ("A5 — Essayages", "Employé", "Clients non rappelés ; rendez-vous oubliés"),
        ("A6 — Préparation livraison", "Employé", "Confusion entre commandes similaires ; retards non anticipés"),
        ("A7 — Livraison et paiement", "Propriétaire", "Suivi des soldes sur papier ; pas de reçu standardisé"),
        ("A8 — Archivage", "Propriétaire", "Réutilisation difficile des anciennes mesures ; pas de recherche rapide"),
    ]
    table_from_rows(
        doc,
        ["Activité", "Acteur", "Difficultés / problèmes"],
        difficultes,
    )
    body(
        doc,
        "Figure 2 [À insérer] : Diagramme BPMN du processus existant avec piscines Client, "
        "Atelier et flux papier / WhatsApp."
    )

    doc.add_heading("2. Problématique", level=2)
    body(
        doc,
        "Alors que le contexte décrit les problèmes vécus par les acteurs du métier, "
        "la problématique traduit la mission de l'étudiant : concevoir et réaliser une "
        "solution informatique capable de supporter un nouveau processus de travail "
        "plus fiable, traçable et adapté au mobile — socle du cahier de charges et "
        "de la conception des chapitres 3 et 4."
    )
    doc.add_heading("2.1 Nouveau processus proposé", level=3)
    body(
        doc,
        "Le nouveau processus repose sur une plateforme SaaS Kayñiawlu qui centralise les "
        "données par organisation (atelier), synchronise mobile et web, et automatise "
        "certaines tâches (numérotation, rappels, statistiques)."
    )
    body(
        doc,
        "L'employé saisit les informations sur le terrain via l'application mobile, même "
        "sans connexion internet : les données sont stockées localement puis synchronisées "
        "automatiquement. Le propriétaire supervise depuis le mobile ou le backoffice web, "
        "affecte les employés aux points de vente, consulte les tableaux de bord et reçoit "
        "des alertes de livraison. Le super administrateur de la plateforme gère les "
        "organisations, les offres commerciales et le monitoring global."
    )
    nouveau = [
        ("N1", "Créer ou retrouver le client dans l'application", "Employé / propriétaire", "Système"),
        ("N2", "Saisir les mesures via fiche paramétrable (types de vêtement)", "Employé", "Système"),
        ("N3", "Créer la commande liée au PDV et à la fiche de mesures", "Employé", "Système"),
        ("N4", "Suivre les statuts (en attente → en cours → essayage → terminée → livrée)", "Équipe", "Système"),
        ("N5", "Recevoir des rappels automatiques de livraison", "Propriétaire", "Système"),
        ("N6", "Générer et partager un reçu PDF", "Propriétaire", "Système"),
        ("N7", "Superviser plusieurs PDV et l'équipe", "Propriétaire", "Système"),
        ("N8", "Administrer la plateforme et les abonnements", "Superadmin", "Système"),
    ]
    table_from_rows(
        doc,
        ["Code", "Activité", "Acteur humain", "Support"],
        nouveau,
    )
    body(
        doc,
        "Figure 3 [À insérer] : Diagramme BPMN du nouveau processus avec interactions "
        "application mobile, backoffice et API."
    )

    doc.add_heading("2.2 Amélioration par rapport au processus existant", level=3)
    bullet(doc, "Réduction des erreurs de saisie grâce à des formulaires structurés et à la validation API.")
    bullet(doc, "Traçabilité complète : chaque commande possède une référence unique (ex. KLM-0237).")
    bullet(doc, "Disponibilité des données en mode hors-ligne sur mobile (SQLite + file de synchronisation).")
    bullet(doc, "Rappels de livraison automatisés (job Laravel planifié + notifications backoffice).")
    bullet(doc, "Vision consolidée multi-PDV pour les ateliers en expansion.")
    bullet(doc, "Modèle économique SaaS avec plans gratuit, pro et entreprise.")

    doc.add_heading("2.3 Objectif général", level=3)
    body(
        doc,
        f"Objectif général : {TITRE_MEMOIRE}."
    )

    doc.add_heading("2.4 Objectifs spécifiques fonctionnels", level=3)
    fonc = [
        ("Propriétaire", "Gérer les points de vente, l'équipe et les abonnements"),
        ("Propriétaire", "Consulter le tableau de bord et les statistiques d'atelier"),
        ("Employé / propriétaire", "CRUD clients avec genre, téléphones multiples, recherche et filtres"),
        ("Employé / propriétaire", "CRUD fiches de mesures (champs JSON par type de vêtement)"),
        ("Employé / propriétaire", "CRUD commandes, statuts, acomptes, dates de livraison"),
        ("Employé / propriétaire", "Recevoir des rappels de livraison"),
        ("Employé / propriétaire", "Générer et partager des reçus PDF"),
        ("Superadmin", "Superviser organisations, utilisateurs et offres SaaS"),
        ("Visiteur", "Consulter la landing page, tarifs et inscription"),
    ]
    table_from_rows(doc, ["Acteur", "Fonctionnalité / service"], fonc)

    doc.add_heading("2.5 Résolution des difficultés par les fonctionnalités", level=3)
    resolution = [
        ("Doublons et historique client dispersé", "F1 — Base clients centralisée multi-utilisateurs"),
        ("Mesures non standardisées", "F2 — Types de vêtement et fiches paramétrables"),
        ("Erreurs sur montants et numérotation", "F3 — Commandes avec référence auto et champs validés"),
        ("Absence de vue globale production", "F4 — Tableau de bord et filtres par statut / PDV"),
        ("Essayages oubliés", "F5 — Rappels livraison et notifications"),
        ("Retards non anticipés", "F5 — Job kalmy:rappels-livraison"),
        ("Pas de reçu standard", "F6 — Export PDF commande"),
        ("Archivage difficile", "F7 — Snapshot immuable des fiches de mesures"),
    ]
    table_from_rows(doc, ["Difficulté", "Fonctionnalité de réponse"], resolution)

    doc.add_heading("2.6 Objectifs spécifiques non fonctionnels", level=3)
    nfunc = [
        ("Architecture", "Monolithique modulaire (API Laravel unique, clients séparés)"),
        ("Interfaces", "Mobile (Flutter) + Web (Next.js 16)"),
        ("Sécurité", "Authentification Sanctum, isolation multi-tenant par organisation"),
        ("Disponibilité", "Mode hors-ligne mobile avec synchronisation"),
        ("Base de données", "PostgreSQL relationnel"),
        ("Coût", "Stack open source (Laravel, Flutter, Next.js)"),
        ("Déploiement", "Conteneurs Docker Compose pour démo / pré-production"),
        ("Qualité", "Tests automatisés PHPUnit (42 tests feature)"),
        ("Scalabilité", "Modèle SaaS par paliers (gratuit, pro, entreprise)"),
    ]
    table_from_rows(doc, ["Contrainte", "Description"], nfunc)

    doc.add_heading("3. Résultats attendus du mémoire", level=2)
    doc.add_heading("3.1 Fonctionnalités réalisées durant le stage", level=3)
    realises = [
        "Authentification Sanctum, inscription « Créer mon atelier », multi-tenant",
        "API GraphQL (lectures) + REST CRUD (écritures)",
        "Application mobile : clients, mesures, commandes, profil, mode offline",
        "Backoffice : portail atelier, superadmin, équipe, types de vêtement, mon compte",
        "Rappels livraison internes + Web Push backoffice",
        "Plans SaaS et structure billing (stub Wave)",
        "Landing page commerciale bilingue",
        "Docker Compose + 42 tests Laravel",
    ]
    for r in realises:
        bullet(doc, r)

    doc.add_heading("3.2 Non fonctionnels réalisés", level=3)
    bullet(doc, "Isolation des données par organisation vérifiée par tests.")
    bullet(doc, "Synchronisation offline mobile (SQLite + queue).")
    bullet(doc, "Design system partagé (couleurs indigo nuit, crème, or).")
    bullet(doc, "Documentation README et RESTE-A-FAIRE.md.")

    doc.add_heading("3.3 Hors périmètre du stage (perspectives)", level=3)
    bullet(doc, "Paiement réel Wave / Orange Money.")
    bullet(doc, "SMS / WhatsApp clients.")
    bullet(doc, "Publication App Store / Play Store.")
    bullet(doc, "Push FCM mobile.")

    doc.add_heading("4. Intérêts du sujet", level=2)
    doc.add_heading("4.1 Pour la société et le secteur", level=3)
    bullet(doc, "Modernisation d'un métier traditionnel à fort impact emploi au Sénégal.")
    bullet(doc, "Réduction des pertes financières liées aux commandes oubliées ou mal suivies.")
    bullet(doc, "Préparation à la digitalisation des paiements mobiles (Wave, OM).")
    bullet(doc, "Enjeu si non réalisé : maintien de pratiques inefficaces et perte de compétitivité face à des outils importés peu adaptés au contexte local.")

    doc.add_heading("4.2 Pour l'étudiant", level=3)
    bullet(doc, "Maîtrise full-stack : Laravel, GraphQL, Flutter, Next.js, PostgreSQL, Docker.")
    bullet(doc, "Expérience produit SaaS multi-tenant et gestion de projet réel.")
    bullet(doc, "Compétences UML/BPMN appliquées à un domaine métier concret.")
    bullet(doc, "Renforcement de l'employabilité sur le marché tech sénégalais et remote.")
    doc.add_page_break()


def add_chapter2(doc: Document) -> None:
    doc.add_heading("Chapitre 2 : Solutions similaires (état de l'art)", level=1)
    body(
        doc,
        "Ce chapitre répond à la question : existe-t-il des solutions similaires à Kayñiawlu "
        "capables de couvrir les fonctionnalités F1 à F8 et les contraintes de qualité définies "
        "en problématique ? Nous présentons dix solutions internationales et régionales, "
        "puis une synthèse comparative."
    )

    solutions = [
        (
            "2.1 TailorPad",
            "Logiciel cloud pour tailleurs et couturiers (Inde / international). Gestion clients, "
            "mesures, commandes, facturation. Interface web, abonnement mensuel. Pas de mode "
            "offline avancé ni focus Afrique de l'Ouest.",
        ),
        (
            "2.2 Stitch Labs / Cin7",
            "ERP mode et retail. Gestion inventaire, commandes, multi-canal. Puissant mais "
            "orienté retail occidental, coût élevé, complexité excessive pour un atelier artisanal.",
        ),
        (
            "2.3 Garmentory / ApparelMagic",
            "Solutions B2B mode (production, wholesale). Adaptées aux usines plus qu'aux ateliers "
            "sur mesure locaux.",
        ),
        (
            "2.4 Booker by Mindbody",
            "Prise de rendez-vous et CRM (beauté / services). Certaines couturières l'utilisent "
            "par détournement ; absence de fiches de mesures couture.",
        ),
        (
            "2.5 Odoo (module Vente / CRM)",
            "ERP open source modulaire. Personnalisable mais nécessite intégrateur, pas de "
            "spécialisation couture sénégalaise native.",
        ),
        (
            "2.6 Zoho Inventory + CRM",
            "Suite SaaS PME. Gestion contacts et commandes génériques ; pas de types boubou / "
            "taille basse.",
        ),
        (
            "2.7 Square Appointments",
            "RDV + paiement (US). Peu répandu au Sénégal, pas de gestion mesures.",
        ),
        (
            "2.8 Excel / Google Sheets (pratique courante)",
            "Gratuit, flexible, mais sans workflow, sans offline mobile structuré, erreurs fréquentes.",
        ),
        (
            "2.9 WhatsApp Business + cahiers (pratique courante au Sénégal)",
            "Communication fluide, photos de mesures, mais aucune structuration, pas de multi-PDV.",
        ),
        (
            "2.10 QuickBooks / Sage (PME)",
            "Comptabilité et facturation. Complément possible mais ne couvre pas le métier couture.",
        ),
    ]
    for title, desc in solutions:
        doc.add_heading(title, level=2)
        body(doc, desc)
        body(
            doc,
            "Analyse : cette solution couvre partiellement ou totalement certaines "
            "fonctionnalités de Kayñiawlu, mais présente des écarts sur la mobilité "
            "hors-ligne, l'adaptation au contexte sénégalais ou le modèle économique "
            "visé. Elle est retenue dans le benchmark pour objectiver le choix de "
            "développer une plateforme sur mesure."
        )

    doc.add_heading("2.11 Tableaux comparatifs", level=2)
    body(doc, "Critères fonctionnels F1–F8 (Oui = couvert, Partiel, Non) :")
    comp_f = [
        ("TailorPad", "Oui", "Oui", "Oui", "Oui", "Partiel", "Oui", "Partiel", "Non"),
        ("Odoo", "Partiel", "Partiel", "Oui", "Oui", "Non", "Oui", "Partiel", "Partiel"),
        ("Zoho CRM", "Partiel", "Non", "Oui", "Partiel", "Non", "Partiel", "Non", "Partiel"),
        ("Booker", "Oui", "Non", "Partiel", "Partiel", "Oui", "Partiel", "Non", "Non"),
        ("WhatsApp + papier", "Partiel", "Partiel", "Partiel", "Non", "Non", "Non", "Non", "Non"),
        ("Kayñiawlu (proposé)", "Oui", "Oui", "Oui", "Oui", "Oui", "Oui", "Oui", "Oui"),
    ]
    table_from_rows(
        doc,
        ["Solution", "F1 Clients", "F2 Mesures", "F3 Commandes", "F4 Dashboard", "F5 Rappels", "F6 PDF", "F7 Multi-PDV", "F8 SaaS admin"],
        comp_f,
    )

    body(doc, "Critères non fonctionnels :")
    comp_nf = [
        ("TailorPad", "Cloud SaaS", "Web", "Non", "Propriétaire", "Cloud", "Non"),
        ("Odoo", "Modulaire", "Web", "Partiel", "Open core", "Les deux", "Partiel"),
        ("Kayñiawlu", "Monolithe API", "Mobile+Web", "Oui", "Open source", "Docker/Cloud", "Oui"),
    ]
    table_from_rows(
        doc,
        ["Solution", "Architecture", "Interface", "Offline", "Licence", "Déploiement", "Contexte Sénégal"],
        comp_nf,
    )

    doc.add_heading("2.12 Conclusion du chapitre", level=2)
    body(
        doc,
        "Aucune solution existante ne combine simultanément : spécialisation couture sur mesure, "
        "application mobile hors-ligne, multi-PDV, modèle SaaS abordable, stack open source et "
        "adaptation au contexte sénégalais (FCFA, types de vêtements locaux, Wave/OM en perspective). "
        "Il est donc nécessaire de concevoir et réaliser Kayñiawlu comme solution dédiée."
    )
    doc.add_page_break()


def add_chapter3(doc: Document) -> None:
    doc.add_heading("Chapitre 3 : Analyse de l'existant et conception de la solution", level=1)

    doc.add_heading("3.1 Partie 1 — Situation existante", level=2)
    doc.add_heading("3.1.1 Diagrammes de cas d'utilisation (situation existante)", level=3)
    body(
        doc,
        "Figure 4 [À insérer] — Couturier propriétaire (processus manuel) : accueillir client, "
        "noter mesures sur papier, enregistrer commande, suivre production, livrer."
    )
    body(
        doc,
        "Figure 5 [À insérer] — Employé : prendre mesures, mettre à jour l'état d'avancement "
        "sur cahier ou WhatsApp."
    )
    body(doc, "Les cas d'utilisation existants mettent en évidence l'absence de système central.")

    doc.add_heading("3.1.2 Contraintes de qualité rappelées", level=3)
    body(doc, "Voir Tableau 4 (Chapitre 1) : sécurité multi-tenant, mobile offline, open source, Docker.")

    doc.add_heading("3.2 Partie 2 — Conception de Kayñiawlu", level=2)
    doc.add_heading("3.2.1 Cas d'utilisation de la future solution", level=3)
    uc = [
        ("UC1", "S'authentifier", "Tous les rôles"),
        ("UC2", "Gérer les clients", "Propriétaire, employé"),
        ("UC3", "Gérer les fiches de mesures", "Propriétaire, employé"),
        ("UC4", "Gérer les commandes", "Propriétaire, employé (PDV assignés)"),
        ("UC5", "Gérer l'équipe et les PDV", "Propriétaire"),
        ("UC6", "Consulter le dashboard", "Propriétaire, superadmin"),
        ("UC7", "Administrer la plateforme", "Superadmin"),
        ("UC8", "Synchroniser en offline", "Système / mobile"),
        ("UC9", "Envoyer rappels livraison", "Système"),
    ]
    table_from_rows(doc, ["Code", "Cas d'utilisation", "Acteurs"], uc)
    body(doc, "Figures 6 à 8 [À insérer] : diagrammes UML par acteur (Draw.io / StarUML).")

    doc.add_heading("3.2.2 Conception backend — modèle de domaine", level=3)
    body(
        doc,
        "Le backend Laravel manipule les entités suivantes, isolées par organisation_id "
        "(multi-tenant logique)."
    )
    entites = [
        ("Organisation", "Atelier tenant : nom, slug, plan, statut abonnement"),
        ("User", "Utilisateur : rôle superadmin | proprietaire | employe"),
        ("PointDeVente", "Boutique ou atelier physique rattaché à l'organisation"),
        ("Client", "Client final : nom, genre, téléphones"),
        ("FicheMesures", "Mesures JSON + snapshot immuable (champs, valeurs)"),
        ("TypeVetement", "Modèle de fiche (boubou, costume…) avec image"),
        ("Commande", "Référence, statut, montants FCFA, livraison_prevue"),
        ("Rappel", "Rappel livraison lié à une commande"),
        ("Offer / Subscription", "Offres SaaS et abonnements"),
    ]
    table_from_rows(doc, ["Entité", "Description"], entites)
    body(doc, "Figure 7 [À insérer] : diagramme de classes UML.")
    body(doc, "Figure 8 [À insérer] : schéma relationnel PostgreSQL.")

    doc.add_heading("3.2.3 Conception frontend", level=3)
    body(
        doc,
        "Application mobile Flutter : navigation par shell (accueil, clients, mesures, commandes, "
        "profil). Écrans principaux : splash, onboarding, login, register, dashboard, formulaires "
        "client/commande/mesure, filtres, centre de notifications."
    )
    body(
        doc,
        "Backoffice Next.js 16 : routes /dashboard (atelier), /dashboard/ateliers (superadmin), "
        "clients, commandes, équipe, types-vetement, mon-compte, abonnement, offres."
    )
    body(
        doc,
        "Landing page : sections hero, fonctionnalités, tarifs, téléchargement app, i18n FR/EN."
    )
    body(doc, "Figures 11 à 13 [À insérer] : maquettes Figma ou captures d'écran.")

    doc.add_heading("3.2.4 Architecture fonctionnelle", level=3)
    body(
        doc,
        "Architecture trois-tiers : (1) présentation — Flutter + Next.js ; (2) logique métier — "
        "Laravel (GraphQL queries, REST CRUD, jobs, policies) ; (3) données — PostgreSQL."
    )
    body(
        doc,
        "Figure 9 [À insérer] : diagramme composants (clients → API → SGBD). "
        "Lecture via GraphQL POST /graphql ; écriture via POST /api/{resource}/save."
    )

    doc.add_heading("3.2.5 Outils de conception", level=3)
    bullet(doc, "Draw.io / diagrams.net — BPMN et UML.")
    bullet(doc, "Figma — design system (Fraunces, Plus Jakarta Sans, palette Kayñiawlu).")
    bullet(doc, "Migrations Laravel — schéma physique itératif.")
    doc.add_page_break()


def add_chapter4(doc: Document) -> None:
    doc.add_heading("Chapitre 4 : Réalisation de la solution", level=1)

    doc.add_heading("4.1 Stack technique", level=2)
    stack = [
        ("Backend", "PHP 8.2+, Laravel 11, Sanctum, rebing/graphql-laravel, PostgreSQL"),
        ("Mobile", "Flutter 3.41, Dart 3.11, SQLite (offline), Provider/Riverpod patterns"),
        ("Backoffice", "Next.js 16, React, Tailwind CSS 4, TypeScript"),
        ("Landing", "Next.js, i18n, composants marketing"),
        ("DevOps", "Docker Compose, Dockerfile backend/backoffice, php artisan queue"),
        ("Tests", "PHPUnit — 42 tests feature (auth, tenant, CRUD, GraphQL)"),
    ]
    table_from_rows(doc, ["Couche", "Technologies"], stack)

    doc.add_heading("4.2 Réalisation backend", level=2)
    doc.add_heading("4.2.1 Structure du dépôt backend", level=3)
    body(
        doc,
        "Le backend Laravel est organisé en couches : Models (Eloquent), Http/Controllers "
        "(CRUDController et contrôleurs métier Auth, Billing, WebPush), GraphQL/Queries "
        "(BaseQuery avec filtres organisation), Policies, Services (WebPushNotificationService), "
        "Jobs et Console/Commands (kalmy:rappels-livraison). Les migrations versionnent "
        "le schéma ; les seeders KalmyDemoSeeder, OfferSeeder et TypeVetementSeeder "
        "alimentent un jeu de démonstration réaliste (Atelier Médina, plusieurs PDV, "
        "clients, commandes KLM-xxxx)."
    )
    doc.add_heading("4.2.2 Base de données PostgreSQL", level=3)
    body(
        doc,
        "La base kalmy contient les tables organisations, users, points_de_vente, clients, "
        "client_telephones, fiches_mesures, types_vetement, commandes, rappels, offers, "
        "subscriptions, web_push_subscriptions, notifications, countries. Les clés étrangères "
        "assurent l'intégrité et le cascade par organisation."
    )
    doc.add_heading("4.2.3 API et sécurité", level=3)
    bullet(doc, "Auth : POST /api/auth/login, register ; tokens Bearer Sanctum.")
    bullet(doc, "Lectures : GraphQL (clients, commandes, fichemesures, pointdeventes, dashboard…).")
    bullet(doc, "Écritures : CRUDController avec hooks — POST /api/{resource}/save, DELETE.")
    bullet(doc, "Policies et scopes : employés limités aux PDV assignés pour les commandes.")
    bullet(doc, "Plans SaaS : limites PDV/utilisateurs (gratuit 1/2, pro 5/10, entreprise illimité).")

    doc.add_heading("4.2.4 Jobs et rappels", level=3)
    body(
        doc,
        "Le job planifié kalmy:rappels-livraison crée des rappels J-7, J-3, J-1 avant "
        "livraison_prevue. Notifications backoffice via polling desktop et Web Push (VAPID, "
        "service worker sw.js)."
    )

    doc.add_heading("4.3 Réalisation mobile Flutter", level=2)
    body(
        doc,
        "L'application mobile vise le couturier en déplacement : prise de mesures chez "
        "le client, création de commande en boutique, consultation des rappels. "
        "Le package lib/ regroupe screens/, services/, data/ et widgets/. "
        "Les comptes de démonstration (ousmane@ateliermedina.sn, binta@ateliermedina.sn) "
        "permettent de valider les scénarios propriétaire et employé limité à un PDV."
    )
    bullet(doc, "kalmy_repository + api_client : couche d'accès API GraphQL/REST.")
    bullet(doc, "local_store + sync_service : cache SQLite, file d'attente offline.")
    bullet(doc, "Écrans métier : clients, mesures, commandes, profil atelier, équipe (admin).")
    bullet(doc, "commande_receipt_service : génération PDF partageable.")
    bullet(doc, "Filtres recherche sur clients et commandes (genre, statut, période).")

    doc.add_heading("4.4 Réalisation backoffice Next.js", level=2)
    bullet(doc, "Authentification et layout dashboard par rôle.")
    bullet(doc, "CRUD clients et commandes web ; lecture mesures ; gestion équipe et PDV.")
    bullet(doc, "Superadmin : ateliers, utilisateurs, offres, monitoring.")
    bullet(doc, "RappelNotificationManager + Web Push.")
    bullet(doc, "Pages Mon compte, types de vêtement (upload / Openverse).")

    doc.add_heading("4.5 Landing page", level=2)
    body(
        doc,
        "Site vitrine présentant la value proposition, les plans tarifaires, les captures produit "
        "et les liens de téléchargement (stores — à activer en production)."
    )

    doc.add_heading("4.6 Déploiement", level=2)
    body(
        doc,
        "Docker Compose à la racine du monorepo lance PostgreSQL, l'API Laravel (migrate, seed, "
        "rappels) et le backoffice. Variables : APP_KEY, DB_*, NEXT_PUBLIC_API_HOST, clés VAPID."
    )
    body(doc, "Figure 10 [À insérer] : diagramme de déploiement Docker.")

    doc.add_heading("4.7 Architecture technique globale", level=2)
    body(
        doc,
        "Figure 10 — Synthèse : Mobile/Backoffice/Landing → HTTPS → Laravel API → PostgreSQL ; "
        "queue worker pour jobs asynchrones ; stockage fichiers (images types vêtement)."
    )
    doc.add_page_break()


def add_chapter5(doc: Document) -> None:
    doc.add_heading("Chapitre 5 : Conclusions et perspectives", level=1)

    doc.add_heading("5.1 Conclusions", level=2)
    doc.add_heading("5.1.1 Résultats atteints", level=3)
    body(
        doc,
        "Les objectifs du mémoire fixés au Chapitre 1 ont été atteints pour l'essentiel : "
        "une plateforme SaaS fonctionnelle couvrant la gestion clients, mesures, commandes, "
        "équipe, multi-PDV, rappels internes, abonnements structurés et interfaces mobile + web + "
        "landing. L'architecture GraphQL/REST, le mode offline et les tests automatisés "
        "valident les contraintes de qualité."
    )

    doc.add_heading("5.1.2 Résultats non atteints", level=3)
    non_atteints = [
        ("Paiement Wave / Orange Money", "Endpoints stub ; intégration API à finaliser"),
        ("Notifications SMS / WhatsApp clients", "Non implémenté ; rappels internes seulement"),
        ("Publication stores iOS / Android", "Builds locaux ; pas de fiches store"),
        ("Push FCM mobile", "Bandeau in-app seulement"),
        ("Déploiement production cloud", "Docker local ; HTTPS prod à configurer"),
    ]
    table_from_rows(doc, ["Objectif", "Explication"], non_atteints)

    doc.add_heading("5.2 Perspectives", level=2)
    bullet(doc, "Finaliser Web Push production (clés VAPID prod, HTTPS, queue worker supervisé).")
    bullet(doc, "Intégrer Wave pour monétiser les plans pro et entreprise.")
    bullet(doc, "Déployer sur VPS / Railway avec CI/CD GitHub Actions.")
    bullet(doc, "Publier sur App Store et Play Store avec captures localisées.")
    bullet(doc, "Espace client final (suivi commande / mesures) — V2.")
    bullet(doc, "Multi-langue wolof, statistiques avancées, export comptable.")
    doc.add_page_break()


def add_bibliography(doc: Document) -> None:
    doc.add_heading("Bibliographie", level=1)
    refs = [
        "NDIAYE S., Guide de rédaction d'un mémoire de master — Génie logiciel / SI, ISI, juin 2026.",
        "Fowler M., Patterns of Enterprise Application Architecture, Addison-Wesley, 2002.",
        "Bass L., Clements P., Kazman R., Software Architecture in Practice, 4e ed., Addison-Wesley, 2021.",
        "Larman C., UML et les Design Patterns, 2e ed., Pearson, 2005.",
        "Richardson L., Ruby S., RESTful Web APIs, O'Reilly, 2013.",
    ]
    for i, r in enumerate(refs, 1):
        body(doc, f"[{i}] {r}")

    doc.add_heading("Webographie", level=1)
    web = [
        "https://laravel.com/docs — Documentation Laravel",
        "https://flutter.dev/docs — Documentation Flutter",
        "https://nextjs.org/docs — Documentation Next.js",
        "https://graphql.org/learn/ — Spécification GraphQL",
        "https://www.postgresql.org/docs/ — Documentation PostgreSQL",
        "https://github.com/legeekado/couture_backend — Dépôt API Kayñiawlu",
        "https://github.com/legeekado/couture_mobile — Dépôt mobile",
        "https://github.com/legeekado/couture_backoffice — Dépôt backoffice",
        "https://github.com/legeekado/couture_landingpage — Dépôt landing page",
    ]
    for w in web:
        bullet(doc, w)


def main() -> None:
    doc = Document()
    set_document_defaults(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_chapter1(doc)
    add_chapter2(doc)
    add_chapter3(doc)
    add_chapter4(doc)
    add_chapter5(doc)
    add_bibliography(doc)
    doc.save(OUTPUT)
    print(f"Mémoire généré : {OUTPUT}")


if __name__ == "__main__":
    main()
