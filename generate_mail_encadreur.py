#!/usr/bin/env python3
"""Génère le mail de proposition de thèmes MBA (rapport d'intervention) au format .docx."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent / "Mail_encadreur_MBA_Jacques_Dembi.docx"

DESTINATAIRE = "M. Diattara"
OBJET = "Proposition de thèmes pour mon rapport d'intervention — MBA Management de projet"
EXPEDITEUR = "Jacques Dembi"
TELEPHONE = "[Téléphone]"
EMAIL = "[Adresse e-mail]"


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(8)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_label_value(doc: Document, label: str, value: str, *, bold_value: bool = False) -> None:
    p = doc.add_paragraph()
    r_label = p.add_run(f"{label} ")
    r_label.bold = True
    r_label.font.name = "Calibri"
    r_label.font.size = Pt(11)
    r_value = p.add_run(value)
    r_value.bold = bold_value
    r_value.font.name = "Calibri"
    r_value.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(10)


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def body_mixed(doc: Document, parts: list[tuple[str, bool]]) -> None:
    p = doc.add_paragraph()
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def heading(doc: Document, text: str, level: int = 2) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.color.rgb = RGBColor(0x17, 0x11, 0x2E)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def bullet(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.bold = True
        r_b.font.name = "Calibri"
        r_b.font.size = Pt(11)
        r = p.add_run(text)
    else:
        r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)


def separator(doc: Document) -> None:
    p = doc.add_paragraph("—" * 40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        r.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def signature(doc: Document) -> None:
    doc.add_paragraph("")
    p = doc.add_paragraph(EXPEDITEUR)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(11)
    for line in (TELEPHONE, EMAIL):
        pl = doc.add_paragraph(line)
        for r in pl.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build() -> Document:
    doc = Document()
    set_document_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("Courrier — Rapport d'intervention MBA")
    rt.bold = True
    rt.font.name = "Calibri"
    rt.font.size = Pt(14)
    rt.font.color.rgb = RGBColor(0x17, 0x11, 0x2E)
    title.paragraph_format.space_after = Pt(18)

    add_label_value(doc, "À :", DESTINATAIRE)
    add_label_value(doc, "De :", EXPEDITEUR)
    add_label_value(doc, "Objet :", OBJET, bold_value=True)

    body(doc, f"Bonjour {DESTINATAIRE},")
    body(
        doc,
        "Je me permets de vous recontacter suite à notre échange téléphonique d'hier soir "
        "concernant mon besoin d'encadrement pour mon MBA en spécialité Management de projet.",
    )
    body_mixed(
        doc,
        [
            ("Pour rappel, je suis ", False),
            ("Jacques Dembi", True),
            (", ingénieur en génie logiciel. J'interviens actuellement chez ", False),
            ("Cauridor", True),
            (
                ", une fintech américaine opérant principalement en Afrique — essentiellement "
                "dans les pays anglophones, mais également dans certains pays francophones. "
                "J'y occupe le poste de ",
                False,
            ),
            ("Tech Lead Reconciliation & Reporting", True),
            (
                ". Mes missions portent notamment sur la lutte contre la fraude, la mise en "
                "place d'outils internes de monitoring et le reporting automatisé connecté "
                "à notre data warehouse.",
                False,
            ),
        ],
    )
    body(
        doc,
        "Dans le cadre de mon MBA, je souhaiterais rédiger un rapport d'intervention et je "
        "me permets de vous soumettre deux pistes de thème, sur lesquelles j'aimerais "
        "recueillir votre avis :",
    )

    separator(doc)

    heading(doc, "Thème 1 (prioritaire) — Kayñiawlu : solution SaaS de gestion pour les ateliers de couture au Sénégal")

    body_mixed(
        doc,
        [
            ("Kayñiawlu", True),
            (" (en wolof : « l'atelier de couture dans votre poche ») est une solution SaaS "
             "que je conçois et développe pour digitaliser la gestion des ateliers de couture "
             "au Sénégal. Ce secteur reste largement informel : les couturiers gèrent encore "
             "leurs clients, mesures et commandes sur papier ou via WhatsApp, ce qui entraîne "
             "des pertes d'information, des retards de livraison et une faible visibilité "
             "sur l'activité.", False),
        ],
    )
    body(doc, "L'objectif du projet est de proposer une plateforme multi-tenant accessible depuis le terrain, couvrant notamment :")

    bullets_theme1 = [
        [
            ("la ", False), ("gestion des clients", True), (" et des ", False),
            ("fiches de mesures", True), (" (boubou, taille basse, costume, etc.) ;", False),
        ],
        [
            ("le ", False), ("suivi des commandes", True),
            (" (statuts, acomptes, dates de livraison, rappels automatiques) ;", False),
        ],
        [
            ("la gestion de ", False), ("plusieurs points de vente", True), (" et de ", False),
            ("l'équipe", True), (" (propriétaire, employés, droits d'accès) ;", False),
        ],
        [
            ("un ", False), ("modèle d'abonnement SaaS", True),
            (" (plans gratuit, pro, entreprise) adapté aux réalités locales ;", False),
        ],
        [
            ("une ", False), ("application mobile Flutter", True),
            (" (usage terrain, mode hors ligne avec synchronisation) ;", False),
        ],
        [
            ("une ", False), ("API Laravel", True),
            (" (PostgreSQL, GraphQL, authentification multi-tenant) ;", False),
        ],
        [
            ("un ", False), ("backoffice web Next.js", True),
            (" pour l'administration et le pilotage de l'activité.", False),
        ],
    ]
    for parts in bullets_theme1:
        p = doc.add_paragraph(style="List Bullet")
        for text, bold in parts:
            r = p.add_run(text)
            r.bold = bold
            r.font.name = "Calibri"
            r.font.size = Pt(11)

    body(
        doc,
        "Le rapport d'intervention porterait sur la conduite de projet de cette solution : "
        "analyse du besoin marché, cadrage fonctionnel, choix d'architecture, gestion des "
        "contraintes (connectivité, adoption utilisateur, monétisation), et retour d'expérience "
        "sur le déploiement d'un produit SaaS en contexte sénégalais.",
    )

    separator(doc)

    heading(doc, "Thème 2 — Guindy Manager : outil interne de gestion intégrée pour une ESN en Afrique de l'Ouest")

    body_mixed(
        doc,
        [
            ("Lors de mon passage en ", False),
            ("Direction technique", True),
            (" chez ", False),
            ("Guindy Technology", True),
            (
                ", société de prestation de services informatiques opérant en sous-région "
                "ouest-africaine, j'ai participé au développement de ",
                False,
            ),
            ("Guindy Manager", True),
            (
                ", un ERP interne centralisant l'ensemble des processus opérationnels "
                "de l'entreprise.",
                False,
            ),
        ],
    )
    body(doc, "Il s'agit d'une application Laravel / PostgreSQL / Redis, structurée autour de plusieurs modules :")

    bullets_theme2 = [
        ("Gestion de projet", " : projets, modules, fonctionnalités, tâches, planification (calendrier), suivi des durées et chronométrage ;"),
        ("Ressources humaines", " : personnel, pointages, demandes d'absence, avances sur salaire ;"),
        ("Relation client & commercial", " : clients, prospects, campagnes marketing, prospection ;"),
        ("Assistance & support", " : tickets d'assistance, rapports d'intervention (PDF), indicateurs de suivi ;"),
        ("Finance & trésorerie", " : caisse, entrées/sorties, dépenses, facturation des abonnements, relances et recouvrement ;"),
        ("Communication", " : campagnes WhatsApp (Twilio), newsletters, intégration Slack, envoi de rapports par e-mail ;"),
        ("Base de connaissances interne", " (« Noyaux Guindy ») et reporting (exports PDF/Excel sur l'ensemble des modules)."),
    ]
    for prefix, rest in bullets_theme2:
        bullet(doc, rest, bold_prefix=prefix)

    body(
        doc,
        "Le rapport d'intervention porterait sur la digitalisation des processus internes d'une ESN : "
        "analyse des besoins métier, conception et déploiement d'un outil sur mesure, gestion du "
        "changement au sein des équipes, et mesure de l'impact sur la productivité et la traçabilité.",
    )

    separator(doc)

    body_mixed(
        doc,
        [
            ("Je privilégie personnellement le ", False),
            ("premier thème (Kayñiawlu)", True),
            (
                ", qui correspond à un projet entrepreneurial en cours et à un enjeu de "
                "transformation digitale d'un secteur informel. Le second thème reste néanmoins "
                "une alternative solide, ancrée dans une expérience concrète de direction technique.",
                False,
            ),
        ],
    )
    body(
        doc,
        "Je reste à votre disposition pour toute information complémentaire si nécessaire, "
        "et je suis également disponible pour une rencontre en présentiel si vous le souhaitez.",
    )
    body(
        doc,
        "Par ailleurs, j'attends de votre part votre recommandation finale sur le choix du thème "
        "à retenir, ainsi qu'un plan de rédaction du rapport d'intervention (structure, calendrier, "
        "livrables attendus).",
    )
    body(
        doc,
        "Dans l'attente de votre retour, je vous prie d'agréer, Monsieur, l'expression de mes "
        "salutations distinguées.",
    )

    signature(doc)
    return doc


def main() -> None:
    doc = build()
    doc.save(OUTPUT)
    print(f"Document généré : {OUTPUT}")


if __name__ == "__main__":
    main()
