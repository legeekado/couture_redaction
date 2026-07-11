#!/usr/bin/env python3
"""Génère la réponse au Dr Diattara (confirmation thème + demandes) au format .docx."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent / "Mail_reponse_encadreur_MBA_Jacques_Dembi.docx"

DESTINATAIRE = "Docteur Diattara"
OBJET = "Re: Proposition de thèmes pour mon rapport d'intervention — MBA Management de projet"
EXPEDITEUR = "Jacques Dembi"
INTITULE = (
    "Transformation digitale des ateliers de couture au Sénégal : "
    "conception, conduite de projet et déploiement d'une solution SaaS multi-tenant, "
    "cas de Kayñiawlu"
)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(8)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_label_value(doc: Document, label: str, value: str, *, bold_value: bool = False) -> None:
    p = doc.add_paragraph()
    r_label = p.add_run(f"{label} ")
    r_label.bold = True
    r_label.font.name = "Calibri"
    r_label.font.size = Pt(11)
    r_value = p.add_run(value)
    r_value.bold = bold_value
    r_value.font.name = "Calibri"
    r_value.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(10)


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def body_mixed(doc: Document, parts: list[tuple[str, bool]]) -> None:
    p = doc.add_paragraph()
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def bullet_mixed(doc: Document, parts: list[tuple[str, bool]]) -> None:
    p = doc.add_paragraph(style="List Bullet")
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def quote_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(f"« {text} »")
    r.bold = True
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x17, 0x11, 0x2E)


def signature(doc: Document) -> None:
    doc.add_paragraph("")
    body(doc, "Bien cordialement,")
    doc.add_paragraph("")
    p = doc.add_paragraph(EXPEDITEUR)
    for r in p.runs:
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def build() -> Document:
    doc = Document()
    set_document_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("Courrier — Rapport d'intervention MBA")
    rt.bold = True
    rt.font.name = "Calibri"
    rt.font.size = Pt(14)
    rt.font.color.rgb = RGBColor(0x17, 0x11, 0x2E)
    title.paragraph_format.space_after = Pt(18)

    add_label_value(doc, "À :", DESTINATAIRE)
    add_label_value(doc, "De :", EXPEDITEUR)
    add_label_value(doc, "Objet :", OBJET, bold_value=True)

    body(doc, f"Bonjour {DESTINATAIRE},")
    body(
        doc,
        "Je vous remercie vivement pour votre retour rapide et pour la qualité de votre analyse.",
    )
    body_mixed(
        doc,
        [
            ("Je confirme donc mon choix pour le ", False),
            ("thème 1", True),
            (", avec l'intitulé principal que vous proposez :", False),
        ],
    )
    quote_block(doc, INTITULE)
    body(
        doc,
        "Je partage tout à fait votre analyse : cet angle donne effectivement plus de hauteur "
        "au sujet en le positionnant comme un projet de conduite de transformation digitale "
        "plutôt qu'un simple développement applicatif, ce qui correspond davantage à l'esprit "
        "d'un MBA en Management de projet.",
    )
    body(
        doc,
        "Afin de démarrer la rédaction dans les meilleures conditions, pourriez-vous m'orienter "
        "sur les points suivants :",
    )

    bullet_mixed(
        doc,
        [
            ("Concernant le plan de rédaction : ", False),
            ("souhaitez-vous m'en proposer une trame", True),
            (", ou ", False),
            ("préférez-vous que je vous soumette moi-même une première proposition", True),
            (" ? Dans ce second cas, verriez-vous un inconvénient à ce que je m'appuie sur les "
             "axes que vous avez évoqués (analyse du besoin, cadrage fonctionnel, choix techniques, "
             "modèle économique, contraintes de terrain, stratégie de déploiement) pour construire "
             "ce plan détaillé ?", False),
        ],
    )
    bullet_mixed(
        doc,
        [
            ("Si possible, ", False),
            ("un ou deux exemples de rapports d'intervention", True),
            (" déjà soutenus, si possible sur des thématiques proches (transformation digitale, "
             "SaaS, conduite de projet), qui pourraient me servir de référence de forme et de "
             "niveau d'exigence ;", False),
        ],
    )
    bullet_mixed(
        doc,
        [
            ("Le cas échéant, ", False),
            ("comment souhaitez-vous que l'on procède pour le suivi de la rédaction", True),
            (" : dois-je vous soumettre l'ensemble du document à la fin, un suivi chapitre par "
             "chapitre au fur et à mesure de l'avancement, ou un découpage par groupes de chapitres "
             "que vous jugeriez pertinent ?", False),
        ],
    )

    body(
        doc,
        "Je reste à votre entière disposition pour toute précision complémentaire, et disponible "
        "pour un échange téléphonique ou en présentiel si cela facilite les choses.",
    )
    body(doc, "Encore merci pour votre accompagnement et vos conseils précieux.")

    signature(doc)
    return doc


def main() -> None:
    doc = build()
    doc.save(OUTPUT)
    print(f"Document généré : {OUTPUT}")


if __name__ == "__main__":
    main()
