#!/usr/bin/env python3
"""Génère le plan détaillé du rapport d'intervention MBA (Kayñiawlu) au format .docx."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt, RGBColor
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent / "Plan_rapport_intervention_MBA_Jacques_Dembi.docx"

AUTEUR = "Jacques Dembi"
ENCADREUR = "Dr Moctar Diattara"
TITRE = (
    "Transformation digitale des ateliers de couture au Sénégal : conception, "
    "conduite de projet et déploiement d'une solution SaaS multi-tenant, "
    "cas de Kayñiawlu"
)

NAVY = RGBColor(0x17, 0x11, 0x2E)
GREY = RGBColor(0x66, 0x66, 0x66)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Proposition de plan détaillé")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Rapport d'intervention — MBA Management de projet")
    r2.italic = True
    r2.font.name = "Calibri"
    r2.font.size = Pt(13)
    r2.font.color.rgb = GREY
    p2.paragraph_format.space_after = Pt(24)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"« {TITRE} »")
    r3.bold = True
    r3.font.name = "Calibri"
    r3.font.size = Pt(13)
    p3.paragraph_format.space_after = Pt(30)

    for label, value in [("Candidat :", AUTEUR), ("Encadreur :", ENCADREUR)]:
        pl = doc.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pl.add_run(f"{label} ")
        rl.bold = True
        rl.font.name = "Calibri"
        rl.font.size = Pt(11)
        rv = pl.add_run(value)
        rv.font.name = "Calibri"
        rv.font.size = Pt(11)

    doc.add_page_break()


def h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True


def h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(12.5)
    r.font.color.rgb = RGBColor(0x2E, 0x7D, 0x5B)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True


def sub(doc: Document, code: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    rc = p.add_run(f"{code}  ")
    rc.bold = True
    rc.font.name = "Calibri"
    rc.font.size = Pt(11)
    rt = p.add_run(text)
    rt.font.name = "Calibri"
    rt.font.size = Pt(11)


def note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = GREY


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(11)


def separator(doc: Document) -> None:
    p = doc.add_paragraph("—" * 40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        r.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


def build() -> Document:
    doc = Document()
    set_document_defaults(doc)
    title_page(doc)

    h1(doc, "Introduction générale")
    sub(doc, "•", "Cadre du rapport (MBA Management de projet, statut d'intervention personnelle sur le projet).")
    sub(doc, "•", "Présentation de l'auteur et de sa légitimité sur le sujet (rôle de concepteur et porteur du projet).")
    sub(doc, "•", "Annonce du plan.")

    separator(doc)

    h1(doc, "Partie 1 — Cadre général du projet")

    h2(doc, "Chapitre 1 — Contexte et problématique")
    sub(doc, "1.1", "Le secteur de la couture au Sénégal : poids économique, informalité, enjeux.")
    sub(doc, "1.2", "Cas d'étude : l'Atelier Kalmy's et son processus de gestion actuel.")
    sub(doc, "1.3", "Diagnostic du processus existant (activités, acteurs, difficultés — modélisation BPMN).")
    sub(doc, "1.4", "Problématique managériale : comment conduire la transformation digitale d'un secteur informel via une solution SaaS ?")
    sub(doc, "1.5", "Enjeux économiques, sociaux et technologiques du projet.")

    h2(doc, "Chapitre 2 — Analyse des besoins")
    sub(doc, "2.1", "Méthodologie de recueil du besoin (observation terrain, entretiens).")
    sub(doc, "2.2", "Besoins par acteur (propriétaire, employé, client final).")
    sub(doc, "2.3", "Contraintes de contexte (connectivité, coût, maturité numérique des utilisateurs).")
    sub(doc, "2.4", "Benchmark des solutions existantes.")
    note(doc, "Réutilisable depuis le chapitre 2 du mémoire M2GL.")
    sub(doc, "2.5", "Positionnement différenciant de Kayñiawlu.")

    separator(doc)

    h1(doc, "Partie 2 — Cadrage et conception du projet")

    h2(doc, "Chapitre 3 — Cadrage du projet")
    sub(doc, "3.1", "Objectifs stratégiques et fonctionnels du projet.")
    sub(doc, "3.2", "Périmètre du projet (in/out scope), livrables attendus.")
    sub(doc, "3.3", "Parties prenantes et gouvernance (matrice RACI).")
    sub(doc, "3.4", "Roadmap produit et jalons.")
    sub(doc, "3.5", "Ressources mobilisées (humaines, techniques, budgétaires).")

    h2(doc, "Chapitre 4 — Conception fonctionnelle et architecture de la solution")
    sub(doc, "4.1", "Conception fonctionnelle : acteurs, cas d'usage, parcours utilisateur.")
    note(doc, "Réutilisable depuis le chapitre 3 du mémoire M2GL.")
    sub(doc, "4.2", "Modèle de données et architecture multi-tenant.")
    sub(doc, "4.3", "Choix techniques et justifications (Laravel/GraphQL, Flutter, Next.js).")
    sub(doc, "4.4", "Contraintes de qualité (sécurité, mode hors-ligne, scalabilité).")

    separator(doc)

    h1(doc, "Partie 3 — Conduite et pilotage du projet")
    note(doc, "Volet le plus « MBA » du rapport : à développer en priorité, avec le plus d'analyse originale.")

    h2(doc, "Chapitre 5 — Gouvernance du projet")
    sub(doc, "5.1", "Organisation et rôles (porteur de projet, développeur, éventuels partenaires).")
    sub(doc, "5.2", "Méthodologie de gestion de projet adoptée (agilité, cycle de vie du produit).")
    sub(doc, "5.3", "Outils de pilotage et de suivi d'avancement.")
    sub(doc, "5.4", "Pilotage de la qualité (tests, revues).")

    h2(doc, "Chapitre 6 — Modèle économique")
    sub(doc, "6.1", "Analyse de marché et segments cibles.")
    sub(doc, "6.2", "Modèle de revenus SaaS (plans gratuit / pro / entreprise).")
    sub(doc, "6.3", "Structure de coûts et hypothèses de rentabilité.")
    sub(doc, "6.4", "Stratégie d'acquisition et de croissance.")

    h2(doc, "Chapitre 7 — Gestion des risques")
    sub(doc, "7.1", "Identification des risques (techniques, marché, adoption, financiers).")
    sub(doc, "7.2", "Matrice probabilité / impact.")
    sub(doc, "7.3", "Plans de mitigation.")
    sub(doc, "7.4", "Suivi des risques en cours de projet.")

    separator(doc)

    h1(doc, "Partie 4 — Déploiement et transformation")

    h2(doc, "Chapitre 8 — Stratégie de déploiement")
    sub(doc, "8.1", "Infrastructure et environnement (Docker, hébergement cible).")
    sub(doc, "8.2", "Déploiement progressif (pilote → généralisation).")
    sub(doc, "8.3", "Stratégie go-to-market (canaux, partenariats avec associations de couturiers).")
    sub(doc, "8.4", "Indicateurs de succès (KPI produit et business).")

    h2(doc, "Chapitre 9 — Conduite du changement")
    sub(doc, "9.1", "Résistance au changement dans un secteur informel : diagnostic.")
    sub(doc, "9.2", "Stratégie d'accompagnement (formation, communication, ambassadeurs terrain).")
    sub(doc, "9.3", "Retours d'adoption utilisateur.")
    sub(doc, "9.4", "Facteurs clés de succès de l'adoption.")

    separator(doc)

    h1(doc, "Partie 5 — Bilan")

    h2(doc, "Chapitre 10 — Enseignements et recommandations")
    sub(doc, "10.1", "Bilan de la conduite de projet (résultats vs objectifs).")
    sub(doc, "10.2", "Enseignements managériaux, techniques et humains.")
    sub(doc, "10.3", "Recommandations pour la suite du projet.")
    sub(doc, "10.4", "Apports du projet pour le candidat MBA.")

    h1(doc, "Conclusion générale")

    h1(doc, "Bibliographie / Webographie")

    h1(doc, "Annexes")
    sub(doc, "•", "Diagrammes BPMN et diagrammes de cas d'utilisation.")
    sub(doc, "•", "Captures d'écran de la solution (backoffice, mobile, landing page).")
    sub(doc, "•", "Extraits de code ou de schéma de base de données si pertinent.")

    return doc


def main() -> None:
    doc = build()
    doc.save(OUTPUT)
    print(f"Document généré : {OUTPUT}")


if __name__ == "__main__":
    main()
